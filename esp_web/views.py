from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.db.models import Value, CharField, Count
from django.db.models.functions import Coalesce, Concat, NullIf, Trim, Lower
from django.db.models import Q
from django.utils.http import url_has_allowed_host_and_scheme
from django.urls import reverse
from django.core.files.base import ContentFile
from django.utils.text import slugify
from django.utils import timezone
from django.utils.safestring import mark_safe
from django.http import JsonResponse, HttpResponse
from django.views.decorators.http import require_http_methods
from django.conf import settings
import json
import os
import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from markdown import markdown
from bs4 import BeautifulSoup
from PIL import Image
from .models import Proyecto, Especificacion, EspecificacionImagen
from .forms import ProyectoForm, EspecificacionForm


def replace_header_placeholders(doc, proyecto, proyecto_nombre=None, solicitante=None, servicio=None, revision="1", fecha=None):
    """
    Reemplaza los placeholders en el encabezado del documento Word con datos del proyecto.
    
    Args:
        doc: Documento Word
        proyecto: Instancia del modelo Proyecto
        proyecto_nombre: Nombre del proyecto personalizado (opcional)
        solicitante: Solicitante personalizado (opcional)
        servicio: Servicio personalizado (opcional)
        revision: Número de revisión (por defecto "1")
        fecha: Fecha personalizada (opcional, formato DD/MM/YYYY)
    """
    # Usar valores personalizados si se proporcionan, sino usar los del proyecto
    proyecto_val = proyecto_nombre if proyecto_nombre is not None else (proyecto.nombre or '')
    solicitante_val = solicitante if solicitante is not None else (proyecto.solicitante or '')
    servicio_val = servicio if servicio is not None else (proyecto.descripcion or proyecto.ubicacion or '')
    fecha_val = fecha if fecha is not None else (proyecto.fecha_creacion.strftime("%d/%m/%Y") if proyecto.fecha_creacion else '')
    
    # Mapeo de placeholders a valores
    replacements = {
        '<<PROYECTO>>': proyecto_val,
        '<<SOLICITANTE>>': solicitante_val,
        '<<SERVICIO>>': servicio_val,
        '<<REV>>': revision,
        '<<REV.>>': revision,  # Por si tiene punto
        '<<FECHA>>': fecha_val,
    }
    
    def replace_in_element(element):
        """Función auxiliar para reemplazar placeholders en un elemento"""
        # Si es un párrafo, trabajar con su texto completo
        if hasattr(element, 'runs'):
            # Primero obtener todo el texto del párrafo
            full_text = ''.join([run.text for run in element.runs])
            
            # Si hay algún placeholder, reemplazar
            if any(ph in full_text for ph in replacements.keys()):
                # Aplicar reemplazos
                new_text = full_text
                for placeholder, value in replacements.items():
                    new_text = new_text.replace(placeholder, value)
                
                # Limpiar todos los runs y agregar el texto reemplazado
                # Preservar el formato del primer run si existe
                if element.runs:
                    # Guardar formato del primer run
                    first_run = element.runs[0]
                    # Limpiar todos los runs
                    for run in element.runs:
                        run.text = ''
                    # Agregar texto con el formato del primer run
                    first_run.text = new_text
                else:
                    element.add_run(new_text)
        elif hasattr(element, 'text'):
            # Para otros elementos con texto directo
            if element.text:
                for placeholder, value in replacements.items():
                    if placeholder in element.text:
                        element.text = element.text.replace(placeholder, value)
    
    # Reemplazar en todas las secciones del documento
    for section in doc.sections:
        # Reemplazar en el encabezado
        header = section.header
        
        # Reemplazar en párrafos del encabezado
        for paragraph in header.paragraphs:
            replace_in_element(paragraph)
        
        # También buscar en tablas del encabezado
        for table in header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_in_element(paragraph)


def _get_especificaciones_accesibles(request):
    qs = (
        Especificacion.objects.filter(
            proyecto__activo=True
        )
        .filter(
            Q(proyecto__publico=True) | Q(proyecto__creado_por=request.user)
        )
        .select_related('proyecto')
        .order_by('proyecto__nombre', '-fecha_creacion')
    )
    especificaciones = []
    for especificacion in qs:
        preview_html = markdown(especificacion.contenido or '', extensions=['extra'])
        especificacion.preview_html = mark_safe(preview_html)
        especificaciones.append(especificacion)
    return especificaciones


def _copiar_especificaciones(user, especificaciones_ids, proyecto_destino):
    copiados = 0
    for spec_id in especificaciones_ids:
        try:
            especificacion = Especificacion.objects.select_related('proyecto').get(
                id=spec_id,
                proyecto__activo=True
            )
        except Especificacion.DoesNotExist:
            continue

        if not (especificacion.proyecto.publico or especificacion.proyecto.creado_por == user):
            continue

        base_titulo = especificacion.titulo
        nuevo_titulo = base_titulo
        contador = 1
        while Especificacion.objects.filter(proyecto=proyecto_destino, titulo=nuevo_titulo).exists():
            if contador == 1:
                nuevo_titulo = f"{base_titulo} (Copia)"
            else:
                nuevo_titulo = f"{base_titulo} (Copia {contador})"
            contador += 1

        nueva_especificacion = Especificacion(
            proyecto=proyecto_destino,
            titulo=nuevo_titulo,
            contenido=especificacion.contenido,
            token_cost=especificacion.token_cost,
            especificacion_tecnica=especificacion.especificacion_tecnica,  # Copiar vinculación si existe
        )

        slug = slugify(nueva_especificacion.titulo) or 'especificacion'
        filename = f"{slug}-{timezone.now():%Y%m%d%H%M%S}.md"
        nueva_especificacion.archivo.save(filename, ContentFile(especificacion.contenido), save=False)
        nueva_especificacion.save()
        copiados += 1
    return copiados


@login_required
def crear_proyecto_view(request):
    """
    Vista para crear un nuevo proyecto
    """
    if request.method == 'POST':
        form = ProyectoForm(request.POST)
        if form.is_valid():
            proyecto = form.save(commit=False)
            if request.user.is_authenticated:
                proyecto.creado_por = request.user
            proyecto.save()
            messages.success(request, f'Proyecto "{proyecto.nombre}" creado exitosamente.')
            # Redirigir a la página principal o a la lista de proyectos
            return redirect('main:proyecto_main')
    else:
        form = ProyectoForm()
    
    return render(request, 'esp_web/crear_proyecto.html', {
        'form': form
    })


@login_required
def lista_proyectos_view(request):
    """
    Vista para listar todos los proyectos existentes
    """
    proyectos = Proyecto.objects.filter(activo=True).order_by('-fecha_creacion')
    
    return render(request, 'esp_web/lista_proyectos.html', {
        'proyectos': proyectos
    })


@login_required
def proyecto_detalle_view(request, proyecto_id):
    """
    Vista de detalle del proyecto seleccionado
    """
    proyecto = get_object_or_404(Proyecto, id=proyecto_id, activo=True)
    es_propietario = proyecto.creado_por == request.user

    # Persistir el proyecto activo en la sesión
    request.session['proyecto_actual_id'] = proyecto.id
    request.session['proyecto_actual_nombre'] = proyecto.nombre

    especificaciones = proyecto.especificaciones.prefetch_related('imagenes').all()
    
    # Inicializar el campo orden si no está establecido (solo si todas tienen orden 0)
    especificaciones_list = list(especificaciones)
    if especificaciones_list and all(spec.orden == 0 for spec in especificaciones_list):
        for i, spec in enumerate(especificaciones_list, start=1):
            spec.orden = i
            spec.save(update_fields=['orden'])
        # Recargar las especificaciones con el nuevo orden
        especificaciones = proyecto.especificaciones.prefetch_related('imagenes').all()
    
    # Obtener ubicaciones del proyecto
    from ubi_web.models import Ubicacion
    ubicaciones = proyecto.ubicaciones.prefetch_related('imagenes').all()
    # Verificar si hay alguna ubicación con PDF generado
    tiene_ubicacion_con_pdf = any(ubicacion.documento_pdf for ubicacion in ubicaciones)

    spec_sort_by = request.GET.get('spec_sort_by', 'titulo')
    spec_order = request.GET.get('spec_order', 'asc')

    valid_spec_sort_fields = ['titulo', 'proyecto', 'usuario']
    if spec_sort_by not in valid_spec_sort_fields:
        spec_sort_by = 'titulo'

    if spec_order not in ['asc', 'desc']:
        spec_order = 'asc'

    especificaciones_accesibles = _get_especificaciones_accesibles(request)
    spec_modal_open = request.GET.get('spec_modal_open') == '1'

    key_map = {
        'titulo': lambda e: (e.titulo or '').lower(),
        'proyecto': lambda e: (e.proyecto.nombre or '').lower(),
        'usuario': lambda e: (
            (
                e.proyecto.creado_por.get_full_name()
                or e.proyecto.creado_por.username
            ).lower()
            if e.proyecto.creado_por
            else ''
        ),
    }

    especificaciones_accesibles.sort(
        key=key_map[spec_sort_by],
        reverse=(spec_order == 'desc')
    )

    return render(request, 'esp_web/proyecto_detalle.html', {
        'proyecto': proyecto,
        'especificaciones': especificaciones,
        'tiene_ubicacion_con_pdf': tiene_ubicacion_con_pdf,
        'ubicaciones': ubicaciones,
        'es_propietario': es_propietario,
        'especificaciones_accesibles': especificaciones_accesibles,
        'spec_sort_by': spec_sort_by,
        'spec_order': spec_order,
        'spec_modal_open': spec_modal_open,
    })


@login_required
def exportar_proyecto_word_view(request, proyecto_id):
    """
    Vista para exportar todas las especificaciones de un proyecto a un documento Word
    """
    proyecto = get_object_or_404(Proyecto, id=proyecto_id, activo=True)
    
    # Verificar permisos
    if not (proyecto.publico or proyecto.creado_por == request.user):
        messages.error(request, 'No tienes permisos para exportar este proyecto.')
        return redirect('main:proyecto_main')
    
    # Obtener valores personalizados del formulario si es POST
    proyecto_nombre = None
    solicitante = None
    servicio = None
    revision = "1"
    fecha = None
    
    if request.method == 'POST':
        proyecto_nombre = request.POST.get('proyecto', proyecto.nombre)
        solicitante = request.POST.get('solicitante', proyecto.solicitante)
        servicio = request.POST.get('servicio', proyecto.descripcion or proyecto.ubicacion)
        revision = request.POST.get('revision', '1')
        fecha = request.POST.get('fecha', proyecto.fecha_creacion.strftime("%d/%m/%Y") if proyecto.fecha_creacion else '')
    
    # Obtener todas las especificaciones ordenadas
    especificaciones = proyecto.especificaciones.prefetch_related('imagenes').all().order_by('orden', '-fecha_creacion')
    
    # Obtener ubicaciones del proyecto
    from ubi_web.models import Ubicacion
    ubicaciones = proyecto.ubicaciones.all()
    
    # Intentar cargar template de Word si existe, sino crear documento nuevo
    template_path = os.path.join(settings.BASE_DIR, 'esp_web', 'templates', 'word_templates', 'template_especificaciones.docx')
    
    # Crear directorio si no existe
    template_dir = os.path.dirname(template_path)
    os.makedirs(template_dir, exist_ok=True)
    
    if os.path.exists(template_path):
        # Cargar el template (conserva encabezados y pies de página)
        doc = Document(template_path)
        # Reemplazar placeholders en el encabezado con datos del proyecto (o valores personalizados)
        replace_header_placeholders(
            doc, 
            proyecto,
            proyecto_nombre=proyecto_nombre,
            solicitante=solicitante,
            servicio=servicio,
            revision=revision,
            fecha=fecha
        )
    else:
        # Crear documento nuevo sin template
        doc = Document()
    
    # Configurar estilos
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Título del documento
    title = doc.add_heading(proyecto.nombre, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Espacio
    
    # Función para agregar texto con formato a un párrafo
    def add_formatted_text(para, elem):
        """Agrega texto con formato (negritas, cursivas, etc.) a un párrafo"""
        if isinstance(elem, str):
            para.add_run(elem)
            return
        
        if not hasattr(elem, 'name') or elem.name is None:
            text = str(elem)
            if text and text.strip():
                para.add_run(text)
            return
        
        tag_name = elem.name.lower() if elem.name else None
        if not tag_name:
            return
        
        # Procesar según el tipo de elemento
        if tag_name in ['strong', 'b']:
            # Negritas
            for child in elem.children:
                if isinstance(child, str):
                    run = para.add_run(child)
                    run.bold = True
                elif hasattr(child, 'name') and child.name:
                    add_formatted_text(para, child)
                else:
                    text = str(child).strip()
                    if text:
                        run = para.add_run(text)
                        run.bold = True
        elif tag_name in ['em', 'i']:
            # Cursivas
            for child in elem.children:
                if isinstance(child, str):
                    run = para.add_run(child)
                    run.italic = True
                elif hasattr(child, 'name') and child.name:
                    add_formatted_text(para, child)
                else:
                    text = str(child).strip()
                    if text:
                        run = para.add_run(text)
                        run.italic = True
        elif tag_name == 'code':
            text = elem.get_text()
            if text:
                run = para.add_run(text)
                run.font.name = 'Courier New'
        elif tag_name == 'a':
            href = elem.get('href', '')
            text = elem.get_text()
            if text:
                run = para.add_run(text if not href else f"{text} ({href})")
                run.font.color.rgb = RGBColor(0, 0, 255)
                run.underline = True
        elif tag_name in ['u']:
            text = elem.get_text()
            if text:
                run = para.add_run(text)
                run.underline = True
        else:
            for child in elem.children:
                add_formatted_text(para, child)
    
    def filtrar_plus_codes(texto):
        """
        Filtra Plus Codes de Google Maps (como "6VF4+2G4") del texto.
        Los Plus Codes tienen formato: letras/números cortos con + seguido de más letras/números.
        """
        if not texto:
            return texto
        
        import re
        # Patrón principal para detectar Plus Codes completos con el símbolo +
        # Formato típico: 4-6 caracteres alfanuméricos + 2-4 caracteres alfanuméricos
        # Ejemplos: "6VF4+2G4", "6VF4+2G", "ABC123+XY"
        patron_plus_code = r'\b[A-Z0-9]{4,6}\+[A-Z0-9]{2,4}\b'
        texto_limpio = re.sub(patron_plus_code, '', texto)
        
        # También filtrar códigos cortos que parecen parte de Plus Codes sin el +
        # Solo si están aislados (rodeados de espacios, comas, puntos, etc.)
        # Ejemplos: "6VF4", "ABC123" cuando están solos y no son parte de direcciones
        def filtrar_codigo_aislado(match):
            codigo = match.group(0)
            # No eliminar si tiene más de 6 caracteres (probablemente no es Plus Code)
            if len(codigo) > 6:
                return codigo
            
            # Obtener contexto alrededor del código
            inicio = max(0, match.start() - 30)
            fin = min(len(texto), match.end() + 30)
            contexto = texto[inicio:fin].lower()
            
            # No eliminar si está cerca de palabras de dirección comunes
            palabras_direccion = ['avenida', 'av.', 'av ', 'calle', 'ruta', 'carretera', 
                                 'km', 'nro', 'número', 'numero', 'dirección', 'direccion',
                                 'barrio', 'zona', 'distrito']
            if any(palabra in contexto for palabra in palabras_direccion):
                return codigo
            
            # No eliminar si está después de "N°", "Nro", "Número", etc.
            if re.search(r'(n[°ºo]|numero|nro|número)\s*' + re.escape(codigo), contexto, re.IGNORECASE):
                return codigo
            
            # Eliminar si es un código corto alfanumérico aislado (probablemente Plus Code)
            return ''
        
        # Buscar códigos de 4-6 caracteres alfanuméricos que estén aislados
        patron_codigo_aislado = r'\b[A-Z0-9]{4,6}\b(?=\s|$|,|\.|;|:|\n)'
        texto_limpio = re.sub(patron_codigo_aislado, filtrar_codigo_aislado, texto_limpio)
        
        # Limpiar espacios múltiples pero preservar estructura de tablas markdown
        # No reemplazar espacios múltiples dentro de líneas de tabla (que empiezan con |)
        lineas = texto_limpio.split('\n')
        lineas_limpias = []
        for linea in lineas:
            # Si es una línea de tabla (contiene |), preservarla tal cual
            if '|' in linea:
                lineas_limpias.append(linea)
            else:
                # Para otras líneas, limpiar espacios múltiples
                linea_limpia = re.sub(r' +', ' ', linea)
                lineas_limpias.append(linea_limpia)
        
        texto_limpio = '\n'.join(lineas_limpias)
        
        # Limpiar saltos de línea múltiples pero preservar al menos uno entre secciones
        texto_limpio = re.sub(r'\n{3,}', '\n\n', texto_limpio)
        
        return texto_limpio.strip()
    
    # Función recursiva para procesar elementos markdown
    def process_markdown_content(contenido_markdown, omitir_titulos=None):
        """Procesa contenido markdown y lo agrega al documento Word
        
        Args:
            contenido_markdown: Contenido en formato markdown
            omitir_titulos: Lista de títulos a omitir (por defecto, None)
        """
        if not contenido_markdown or not contenido_markdown.strip():
            return
        
        # Convertir markdown a HTML y luego procesar
        contenido_html = markdown(contenido_markdown, extensions=['extra'])
        soup = BeautifulSoup(contenido_html, 'html.parser')
        
        # Lista de títulos a omitir (si no se proporciona, usar lista vacía)
        if omitir_titulos is None:
            omitir_titulos = []
        
        # Función recursiva para procesar elementos
        def process_element(elem):
            if not hasattr(elem, 'name') or elem.name is None:
                text = str(elem).strip()
                if text and text not in ['\n', '\r', '\t', '']:
                    clean_text = ' '.join(text.split())
                    if clean_text:
                        para = doc.add_paragraph()
                        para.add_run(clean_text)
                return
            
            tag_name = elem.name.lower() if elem.name else None
            if not tag_name:
                return
            
            # Detectar headings (h1-h6)
            if tag_name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                # Obtener el texto del heading
                heading_text = elem.get_text().strip()
                # Verificar si este título debe omitirse
                if heading_text not in omitir_titulos:
                    # Solo agregar el heading si NO está en la lista de omitir
                    level = int(tag_name[1])
                    # Limitar el nivel máximo a 2 (heading 2)
                    if level > 2:
                        level = 2
                    heading = doc.add_heading(level=level)
                    for child in elem.children:
                        add_formatted_text(heading, child)
                # No hacer return aquí - permitir que el procesamiento continúe
                # Los elementos siguientes se procesarán normalmente en el bucle principal
                return
                
            
            # Detectar párrafos
            elif tag_name == 'p':
                # Verificar si este párrafo ya fue procesado (marcado con _processed)
                if elem.get('_processed'):
                    return
                
                # Verificar si este párrafo contiene una imagen
                img_in_p = elem.find('img')
                if img_in_p:
                    # Si hay una imagen en el párrafo, procesarla primero
                    process_element(img_in_p)
                    # Luego procesar el resto del contenido del párrafo (si hay texto después de la imagen)
                    for child in elem.children:
                        if child != img_in_p and child.name != 'img':
                            if isinstance(child, str):
                                text = child.strip()
                                if text:
                                    para = doc.add_paragraph()
                                    para.add_run(text)
                            elif hasattr(child, 'name') and child.name:
                                # Si es texto en cursiva que parece un pie de figura, procesarlo después
                                if child.name.lower() in ['em', 'i']:
                                    text = child.get_text().strip()
                                    if text and ('Figura' in text or 'figura' in text):
                                        caption_para = doc.add_paragraph(text)
                                        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                        caption_para.style.font.size = Pt(9)
                                        caption_para.style.font.italic = True
                                    else:
                                        para = doc.add_paragraph()
                                        add_formatted_text(para, child)
                                else:
                                    para = doc.add_paragraph()
                                    add_formatted_text(para, child)
                else:
                    # Verificar si este párrafo es un pie de figura (texto en cursiva que contiene "Figura")
                    em_or_i = elem.find(['em', 'i'])
                    if em_or_i:
                        text = em_or_i.get_text().strip()
                        if text and ('Figura' in text or 'figura' in text):
                            # Este es un pie de figura, verificar si hay una imagen antes
                            # Buscar el elemento anterior que pueda ser una imagen
                            prev_elem = elem.find_previous_sibling(['p', 'img'])
                            if prev_elem:
                                if prev_elem.name == 'img':
                                    # Ya se procesó la imagen, solo agregar el pie de figura
                                    caption_para = doc.add_paragraph(text)
                                    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    caption_para.style.font.size = Pt(9)
                                    caption_para.style.font.italic = True
                                    return
                                elif prev_elem.name == 'p':
                                    prev_img = prev_elem.find('img')
                                    if prev_img:
                                        # Ya se procesó la imagen, solo agregar el pie de figura
                                        caption_para = doc.add_paragraph(text)
                                        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                        caption_para.style.font.size = Pt(9)
                                        caption_para.style.font.italic = True
                                        return
                    
                    # Procesamiento normal del párrafo
                    para = doc.add_paragraph()
                    for child in elem.children:
                        add_formatted_text(para, child)
            
            # Detectar listas
            elif tag_name == 'ul':
                for li in elem.find_all('li', recursive=False):
                    para = doc.add_paragraph(style='List Bullet')
                    for child in li.children:
                        add_formatted_text(para, child)
            
            elif tag_name == 'ol':
                for li in elem.find_all('li', recursive=False):
                    para = doc.add_paragraph(style='List Number')
                    for child in li.children:
                        add_formatted_text(para, child)
            
            elif tag_name == 'li':
                para = doc.add_paragraph(style='List Bullet')
                for child in elem.children:
                    add_formatted_text(para, child)
            
            elif tag_name == 'br':
                doc.add_paragraph()
            
            # Detectar tablas
            elif tag_name == 'table':
                # Crear tabla en Word
                rows = elem.find_all('tr', recursive=False)
                if rows:
                    num_rows = len(rows)
                    # Contar columnas de la primera fila
                    first_row = rows[0]
                    num_cols = len(first_row.find_all(['td', 'th'], recursive=False))
                    if num_cols > 0:
                        table = doc.add_table(rows=num_rows, cols=num_cols)
                        table.style = 'Light Grid Accent 1'
                        
                        # Llenar la tabla
                        for row_idx, row in enumerate(rows):
                            cells = row.find_all(['td', 'th'], recursive=False)
                            for col_idx, cell in enumerate(cells):
                                if col_idx < num_cols:
                                    word_cell = table.rows[row_idx].cells[col_idx]
                                    # Limpiar párrafos existentes
                                    word_cell.text = ''
                                    # Agregar contenido con formato
                                    for child in cell.children:
                                        add_formatted_text(word_cell.paragraphs[0], child)
                        
                        # Eliminar bordes de la tabla
                        for row in table.rows:
                            for cell in row.cells:
                                tcPr = cell._element.tcPr
                                if tcPr is None:
                                    tcPr = OxmlElement('w:tcPr')
                                    cell._element.append(tcPr)
                                
                                tcBorders = tcPr.find(qn('w:tcBorders'))
                                if tcBorders is None:
                                    tcBorders = OxmlElement('w:tcBorders')
                                    tcPr.append(tcBorders)
                                
                                for border_name in ['top', 'left', 'bottom', 'right']:
                                    border = tcBorders.find(qn(f'w:{border_name}'))
                                    if border is None:
                                        border = OxmlElement(f'w:{border_name}')
                                        tcBorders.append(border)
                                    border.set(qn('w:val'), 'nil')
                                    border.set(qn('w:sz'), '0')
                                    border.set(qn('w:space'), '0')
            
            # Detectar imágenes
            elif tag_name == 'img':
                src = elem.get('src', '')
                alt = elem.get('alt', '')
                if src:
                    # Intentar obtener la ruta del archivo desde la URL
                    try:
                        # Si es una URL relativa (media), convertir a ruta de archivo
                        if src.startswith('/media/'):
                            # Remover /media/ del inicio
                            media_path = src.replace('/media/', '')
                            imagen_path = os.path.join(settings.MEDIA_ROOT, media_path)
                        elif src.startswith('media/'):
                            imagen_path = os.path.join(settings.MEDIA_ROOT, src.replace('media/', ''))
                        else:
                            # Intentar como ruta absoluta o relativa
                            imagen_path = src
                            if not os.path.isabs(imagen_path):
                                imagen_path = os.path.join(settings.MEDIA_ROOT, imagen_path)
                        
                        if os.path.exists(imagen_path):
                            para = doc.add_paragraph()
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
                            img = Image.open(imagen_path)
                            max_width_inches = 5.0
                            max_height_inches = 4.0
                            
                            img_width_px = img.width
                            img_height_px = img.height
                            
                            try:
                                dpi = img.info.get('dpi', (96, 96))[0]
                            except:
                                dpi = 96
                            
                            width_inches = img_width_px / dpi
                            height_inches = img_height_px / dpi
                            
                            if width_inches > max_width_inches or height_inches > max_height_inches:
                                ratio = min(max_width_inches / width_inches, max_height_inches / height_inches)
                                width_inches *= ratio
                                height_inches *= ratio
                            
                            run = para.add_run()
                            run.add_picture(imagen_path, width=Inches(width_inches))
                            
                            # Buscar si hay un párrafo siguiente con texto en cursiva que contenga "Figura"
                            # Primero buscar en el mismo párrafo (si la imagen está dentro de un párrafo)
                            parent_p = elem.find_parent('p')
                            if parent_p:
                                # Buscar texto en cursiva después de la imagen en el mismo párrafo
                                em_or_i_after = None
                                for sibling in elem.next_siblings:
                                    if hasattr(sibling, 'name'):
                                        if sibling.name in ['em', 'i']:
                                            em_or_i_after = sibling
                                            break
                                        elif sibling.name == 'p':
                                            break
                                
                                if em_or_i_after:
                                    caption_text = em_or_i_after.get_text().strip()
                                    if caption_text and ('Figura' in caption_text or 'figura' in caption_text):
                                        # Agregar el pie de figura después de la imagen
                                        caption_para = doc.add_paragraph(caption_text)
                                        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                        caption_para.style.font.size = Pt(9)
                                        caption_para.style.font.italic = True
                                        return
                                
                                # Si no está en el mismo párrafo, buscar en el siguiente párrafo hermano
                                next_p = parent_p.find_next_sibling('p')
                                if next_p:
                                    em_or_i = next_p.find(['em', 'i'])
                                    if em_or_i:
                                        caption_text = em_or_i.get_text().strip()
                                        if caption_text and ('Figura' in caption_text or 'figura' in caption_text):
                                            # Agregar el pie de figura después de la imagen
                                            caption_para = doc.add_paragraph(caption_text)
                                            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                            caption_para.style.font.size = Pt(9)
                                            caption_para.style.font.italic = True
                                            # Marcar este elemento para que no se procese de nuevo
                                            next_p['_processed'] = True
                                            return
                            else:
                                # Si la imagen no está dentro de un párrafo, buscar el siguiente elemento hermano
                                next_elem = elem.find_next_sibling('p')
                                if next_elem:
                                    em_or_i = next_elem.find(['em', 'i'])
                                    if em_or_i:
                                        caption_text = em_or_i.get_text().strip()
                                        if caption_text and ('Figura' in caption_text or 'figura' in caption_text):
                                            # Agregar el pie de figura después de la imagen
                                            caption_para = doc.add_paragraph(caption_text)
                                            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                            caption_para.style.font.size = Pt(9)
                                            caption_para.style.font.italic = True
                                            # Marcar este elemento para que no se procese de nuevo
                                            next_elem['_processed'] = True
                                            return
                            
                            # Si no hay pie de figura siguiente, usar el texto alternativo si existe
                            if alt:
                                caption_para = doc.add_paragraph(alt)
                                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                caption_para.style.font.size = Pt(9)
                                caption_para.style.font.italic = True
                    except Exception as e:
                        # Si hay error, agregar texto alternativo
                        para = doc.add_paragraph()
                        para.add_run(f'[Imagen: {alt if alt else "No disponible"}]')
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            else:
                if hasattr(elem, 'children') and list(elem.children):
                    for child in elem.children:
                        process_element(child)
                else:
                    text = elem.get_text().strip()
                    if text:
                        para = doc.add_paragraph()
                        add_formatted_text(para, elem)
        
        # Procesar todos los elementos principales del HTML en orden
        # Procesar todos los hijos directos del soup para mantener el orden
        for element in soup.children:
            if hasattr(element, 'name') and element.name:
                process_element(element)
            elif isinstance(element, str) and element.strip():
                # Procesar texto suelto
                para = doc.add_paragraph()
                para.add_run(element.strip())
    
    # Agregar contenido de ubicación al principio si existe
    if ubicaciones.exists():
        ubicacion = ubicaciones.first()  # Tomar la primera ubicación
        if ubicacion.contenido:
            # Título de sección de ubicación
            ubicacion_heading = doc.add_heading("Ubicación del Sitio", level=2)
            
            # Filtrar Plus Codes del contenido antes de procesarlo
            contenido_limpio = filtrar_plus_codes(ubicacion.contenido)
            
            # Verificar si hay imágenes en el contenido markdown limpio
            # Si hay imágenes en el markdown, asumimos que la imagen del mapa ya está incluida
            imagen_en_markdown = False
            if contenido_limpio:
                contenido_html_temp = markdown(contenido_limpio, extensions=['extra'])
                soup_temp = BeautifulSoup(contenido_html_temp, 'html.parser')
                imagenes_en_markdown = soup_temp.find_all('img')
                if imagenes_en_markdown:
                    # Si hay al menos una imagen en el markdown, asumimos que es la del mapa
                    imagen_en_markdown = True
            
            # Procesar el contenido markdown de la ubicación, omitiendo títulos específicos
            process_markdown_content(
                contenido_limpio, 
                omitir_titulos=['Coordenadas del Sitio', 'Descripción de Acceso']
            )
            
            # Agregar imagen del mapa solo si existe y NO está ya en el contenido markdown
            # (si el markdown ya tiene imágenes, no agregar la imagen manualmente)
            if ubicacion.mapa_imagen and ubicacion.mapa_imagen.name and not imagen_en_markdown:
                try:
                    mapa_path = ubicacion.mapa_imagen.path
                    if os.path.exists(mapa_path):
                        doc.add_paragraph()
                        para = doc.add_paragraph()
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        img = Image.open(mapa_path)
                        max_width_inches = 5.0
                        max_height_inches = 4.0
                        
                        img_width_px = img.width
                        img_height_px = img.height
                        
                        try:
                            dpi = img.info.get('dpi', (96, 96))[0]
                        except:
                            dpi = 96
                        
                        width_inches = img_width_px / dpi
                        height_inches = img_height_px / dpi
                        
                        if width_inches > max_width_inches or height_inches > max_height_inches:
                            ratio = min(max_width_inches / width_inches, max_height_inches / height_inches)
                            width_inches *= ratio
                            height_inches *= ratio
                        
                        run = para.add_run()
                        run.add_picture(mapa_path, width=Inches(width_inches))
                        
                        # Agregar pie de figura
                        caption_para = doc.add_paragraph(f"Mapa satelital del sitio {ubicacion.nombre}")
                        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        caption_para.style.font.size = Pt(9)
                        caption_para.style.font.italic = True
                except Exception as e:
                    # Si hay error al agregar la imagen, continuar sin ella
                    pass
            
            # Espacio después de la ubicación
            doc.add_paragraph()
            doc.add_paragraph('─' * 50)
            doc.add_paragraph()
    
    # Agregar cada especificación
    for especificacion in especificaciones:
        # Procesar el contenido markdown de la especificación usando la función común
        process_markdown_content(especificacion.contenido)
        
        # Agregar imágenes si existen
        imagenes = especificacion.imagenes.all()
        if imagenes:
            doc.add_paragraph()
            
            num_imagenes = imagenes.count()
            num_filas = (num_imagenes + 1) // 2
            
            # Crear tabla con 2 columnas sin bordes
            table = doc.add_table(rows=num_filas, cols=2)
            
            # Eliminar bordes de la tabla
            for row in table.rows:
                for cell in row.cells:
                    tcPr = cell._element.tcPr
                    if tcPr is None:
                        tcPr = OxmlElement('w:tcPr')
                        cell._element.append(tcPr)
                    
                    tcBorders = tcPr.find(qn('w:tcBorders'))
                    if tcBorders is None:
                        tcBorders = OxmlElement('w:tcBorders')
                        tcPr.append(tcBorders)
                    
                    for border_name in ['top', 'left', 'bottom', 'right']:
                        border = tcBorders.find(qn(f'w:{border_name}'))
                        if border is None:
                            border = OxmlElement(f'w:{border_name}')
                            tcBorders.append(border)
                        border.set(qn('w:val'), 'nil')
                        border.set(qn('w:sz'), '0')
                        border.set(qn('w:space'), '0')
            
            # Llenar la tabla con las imágenes
            imagen_index = 0
            for row in table.rows:
                for cell in row.cells:
                    if imagen_index < num_imagenes:
                        imagen = imagenes[imagen_index]
                        imagen_path = imagen.imagen.path
                        
                        if os.path.exists(imagen_path):
                            try:
                                paragraph = cell.paragraphs[0]
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                
                                img = Image.open(imagen_path)
                                max_width_inches = 3.0
                                max_height_inches = 3.0
                                
                                img_width_px = img.width
                                img_height_px = img.height
                                
                                try:
                                    dpi = img.info.get('dpi', (96, 96))[0]
                                except:
                                    dpi = 96
                                
                                width_inches = img_width_px / dpi
                                height_inches = img_height_px / dpi
                                
                                if width_inches > max_width_inches or height_inches > max_height_inches:
                                    ratio = min(max_width_inches / width_inches, max_height_inches / height_inches)
                                    width_inches *= ratio
                                    height_inches *= ratio
                                
                                run = paragraph.add_run()
                                run.add_picture(imagen_path, width=Inches(width_inches))
                                
                                if imagen.descripcion:
                                    desc_para = cell.add_paragraph(imagen.descripcion)
                                    desc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    desc_para.style.font.size = Pt(9)
                            except Exception as e:
                                error_para = cell.paragraphs[0]
                                error_para.add_run('[Error al cargar imagen]')
                                error_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        imagen_index += 1
            
            doc.add_paragraph()
        
        # Espacio entre especificaciones
        doc.add_paragraph()
        doc.add_paragraph('─' * 50)
        doc.add_paragraph()
    
    # Preparar la respuesta
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    # Usar nombre personalizado si está disponible
    nombre_archivo = proyecto_nombre if proyecto_nombre else proyecto.nombre
    filename = f'{slugify(nombre_archivo)}_especificaciones.docx'
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    
    # Guardar el documento en memoria
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    response.write(buffer.read())
    
    return response


@login_required
def ingresar_proyecto_view(request, proyecto_id):
    """
    Vincula el proyecto activo y redirige al flujo de nuevo pliego
    """
    proyecto = get_object_or_404(Proyecto, id=proyecto_id, activo=True)

    if proyecto.creado_por != request.user:
        messages.error(request, 'Solo puedes crear especificaciones en tus propios proyectos.')
        return redirect('main:proyecto_main')

    request.session['proyecto_actual_id'] = proyecto.id
    request.session['proyecto_actual_nombre'] = proyecto.nombre

    intended = request.GET.get('next')
    if intended == 'detalle':
        return redirect('esp_web:proyecto_detalle', proyecto_id=proyecto.id)

    return redirect(reverse('nuevo_pliego_view'))


@login_required
def seleccionar_proyecto_view(request, proyecto_id):
    """
    Vista para seleccionar un proyecto existente
    """
    proyecto = get_object_or_404(Proyecto, id=proyecto_id, activo=True)
    
    # Guardar el proyecto seleccionado en la sesión
    request.session['proyecto_actual_id'] = proyecto.id
    request.session['proyecto_actual_nombre'] = proyecto.nombre
    
    next_url = request.GET.get('next')
    if next_url:
        # Si la URL no tiene esquema, agregarlo para la validación
        if not next_url.startswith(('http://', 'https://')):
            # Construir URL completa con el host actual
            scheme = 'https' if request.is_secure() else 'http'
            full_next_url = f"{scheme}://{request.get_host()}{next_url}"
        else:
            full_next_url = next_url
        
        if url_has_allowed_host_and_scheme(full_next_url, allowed_hosts={request.get_host()}):
            # Redirigir a la URL original (sin esquema) si es relativa
            if not next_url.startswith(('http://', 'https://')):
                return redirect(next_url)
            return redirect(full_next_url)

    referer = request.META.get('HTTP_REFERER')
    if referer and url_has_allowed_host_and_scheme(referer, allowed_hosts={request.get_host()}):
        return redirect(referer)

    return redirect('main:proyecto_main')


@login_required
def editar_especificacion_view(request, especificacion_id):
    especificacion = get_object_or_404(Especificacion, id=especificacion_id, proyecto__activo=True)
    proyecto = especificacion.proyecto

    if proyecto.creado_por != request.user:
        messages.error(request, 'Solo puedes editar especificaciones de tus proyectos.')
        return redirect('esp_web:proyecto_detalle', proyecto_id=proyecto.id)

    if request.method == 'POST':
        form = EspecificacionForm(request.POST, instance=especificacion)
        if form.is_valid():
            especificacion = form.save(commit=False)
            if especificacion.archivo:
                especificacion.archivo.delete(save=False)
            slug = slugify(especificacion.titulo) or 'especificacion'
            filename = f"{slug}-{timezone.now().strftime('%Y%m%d%H%M%S')}.md"
            especificacion.archivo.save(filename, ContentFile(especificacion.contenido), save=False)
            especificacion.save()
            messages.success(request, 'Especificación actualizada correctamente.')
            return redirect('esp_web:proyecto_detalle', proyecto_id=proyecto.id)
    else:
        form = EspecificacionForm(instance=especificacion)

    return render(request, 'esp_web/editar_especificacion.html', {
        'form': form,
        'especificacion': especificacion,
        'proyecto': proyecto,
    })


@login_required
def eliminar_especificacion_view(request, especificacion_id):
    especificacion = get_object_or_404(Especificacion, id=especificacion_id, proyecto__activo=True)
    proyecto = especificacion.proyecto

    if proyecto.creado_por != request.user:
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest' or request.content_type == 'application/json':
            return JsonResponse({'error': 'Solo puedes eliminar especificaciones de tus proyectos.'}, status=403)
        messages.error(request, 'Solo puedes eliminar especificaciones de tus proyectos.')
        return redirect('esp_web:proyecto_detalle', proyecto_id=proyecto.id)

    if request.method == 'POST':
        if especificacion.archivo:
            especificacion.archivo.delete(save=False)
        especificacion.delete()
        
        # Si es una petición AJAX, devolver JSON
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest' or request.content_type == 'application/json':
            return JsonResponse({
                'success': True,
                'message': 'Especificación eliminada correctamente.'
            })
        
        messages.success(request, 'Especificación eliminada correctamente.')
        return redirect('esp_web:proyecto_detalle', proyecto_id=proyecto.id)

    # Si es GET, mostrar la página de confirmación (para compatibilidad)
    return render(request, 'esp_web/eliminar_especificacion.html', {
        'especificacion': especificacion,
        'proyecto': proyecto,
    })


@login_required
def eliminar_proyecto_view(request, proyecto_id):
    """
    Vista para eliminar (desactivar) un proyecto
    """
    proyecto = get_object_or_404(Proyecto, id=proyecto_id, activo=True)

    if proyecto.creado_por != request.user:
        messages.error(request, 'Solo puedes eliminar proyectos que pertenecen a tu cuenta.')
        return redirect('main:proyecto_main')
    
    if request.method == 'POST':
        proyecto.activo = False
        proyecto.save()
        messages.success(request, f'Proyecto "{proyecto.nombre}" eliminado exitosamente.')
        return redirect('main:proyecto_main')
    
    return render(request, 'esp_web/eliminar_proyecto.html', {
        'proyecto': proyecto
    })


@login_required
def especificaciones_disponibles_view(request):
    sort_by = request.GET.get('sort_by', 'nombre')
    order = request.GET.get('order', 'asc')
    dest_id = request.GET.get('dest')
    dest_project = None

    valid_sort_fields = ['nombre', 'proyecto', 'fecha']
    if sort_by not in valid_sort_fields:
        sort_by = 'nombre'

    if order not in ['asc', 'desc']:
        order = 'asc'

    especificaciones = _get_especificaciones_accesibles(request)

    def _safe_lower(value):
        return (value or '').lower()

    fallback_date = timezone.now()

    key_map = {
        'nombre': lambda e: _safe_lower(getattr(e, 'titulo', '')),
        'proyecto': lambda e: _safe_lower(getattr(e.proyecto, 'nombre', '') if getattr(e, 'proyecto', None) else ''),
        'fecha': lambda e: getattr(e, 'fecha_creacion', None) or fallback_date,
    }

    especificaciones.sort(key=key_map[sort_by], reverse=(order == 'desc'))

    if dest_id:
        try:
            dest_project = Proyecto.objects.get(id=dest_id, activo=True, creado_por=request.user)
        except Proyecto.DoesNotExist:
            dest_project = None

    mis_proyectos = Proyecto.objects.filter(activo=True, creado_por=request.user).order_by('nombre')
    return render(request, 'esp_web/especificaciones_disponibles.html', {
        'especificaciones': especificaciones,
        'mis_proyectos': mis_proyectos,
        'sort_by': sort_by,
        'order': order,
        'dest_id': dest_id,
        'dest_project': dest_project,
    })


@login_required
def ver_especificacion_view(request, especificacion_id):
    especificacion = get_object_or_404(
        Especificacion.objects.select_related('proyecto'),
        id=especificacion_id,
        proyecto__activo=True
    )

    if not (especificacion.proyecto.publico or especificacion.proyecto.creado_por == request.user):
        messages.error(request, 'No tienes permisos para ver esta especificación.')
        return redirect('main:proyecto_main')

    origin = request.GET.get('from')
    if origin not in ('disponibles', 'proyecto'):
        origin = 'proyecto'

    dest_id = request.GET.get('dest')
    dest_project = None
    if origin == 'disponibles' and dest_id:
        try:
            dest_project = Proyecto.objects.get(id=dest_id, activo=True, creado_por=request.user)
        except Proyecto.DoesNotExist:
            dest_project = None

    preview_html = markdown(especificacion.contenido or '', extensions=['extra'])
    preview_html = mark_safe(preview_html)

    return render(request, 'esp_web/ver_especificacion.html', {
        'especificacion': especificacion,
        'preview_html': preview_html,
        'breadcrumbs_origin': origin,
        'dest_project': dest_project,
    })


@login_required
def copiar_especificacion_view(request, especificacion_id):
    if request.method != 'POST':
        messages.error(request, 'Acción inválida al intentar copiar la especificación.')
        return redirect('main:proyecto_main')

    proyecto_destino_id = request.POST.get('proyecto_destino')
    if not proyecto_destino_id:
        messages.error(request, 'Debes seleccionar un proyecto destino.')
        return redirect('main:proyecto_main')

    try:
        proyecto_destino = Proyecto.objects.get(
            id=proyecto_destino_id,
            activo=True,
            creado_por=request.user
        )
    except Proyecto.DoesNotExist:
        messages.error(request, 'Solo puedes copiar especificaciones a tus propios proyectos activos.')
        return redirect('main:proyecto_main')

    copiados = _copiar_especificaciones(request.user, [str(especificacion_id)], proyecto_destino)
    if copiados:
        messages.success(
            request,
            f'Se copió la especificación al proyecto "{proyecto_destino.nombre}".'
        )
    else:
        messages.warning(request, 'No se copiaron especificaciones. Verifica la selección y tus permisos.')

    next_url = request.POST.get('next')
    if next_url and url_has_allowed_host_and_scheme(next_url, allowed_hosts={request.get_host()}):
        return redirect(next_url)
    return redirect('esp_web:proyecto_detalle', proyecto_destino.id)


@login_required
def copiar_especificaciones_view(request):
    if request.method != 'POST':
        messages.error(request, 'Acción inválida al intentar copiar especificaciones.')
        return redirect('main:proyecto_main')

    especificaciones_ids = request.POST.getlist('especificaciones')
    if not especificaciones_ids:
        messages.error(request, 'Selecciona al menos una especificación.')
        return redirect('main:proyecto_main')
    especificaciones_ids = list(dict.fromkeys(str(e) for e in especificaciones_ids))

    proyecto_destino_id = request.POST.get('proyecto_destino')
    if not proyecto_destino_id:
        messages.error(request, 'Debes seleccionar un proyecto destino.')
        return redirect('main:proyecto_main')

    try:
        proyecto_destino = Proyecto.objects.get(
            id=proyecto_destino_id,
            activo=True,
            creado_por=request.user
        )
    except Proyecto.DoesNotExist:
        messages.error(request, 'Solo puedes copiar especificaciones a tus propios proyectos activos.')
        return redirect('main:proyecto_main')

    copiados = _copiar_especificaciones(request.user, especificaciones_ids, proyecto_destino)

    if copiados:
        messages.success(
            request,
            f'Se copiaron {copiados} especificación(es) al proyecto "{proyecto_destino.nombre}".'
        )
    else:
        messages.warning(request, 'No se copiaron especificaciones. Verifica la selección y tus permisos.')

    return redirect('main:proyecto_main')


@login_required
def editar_proyecto_view(request, proyecto_id):
    """
    Vista para editar un proyecto existente
    """
    proyecto = get_object_or_404(Proyecto, id=proyecto_id, activo=True)
    
    if proyecto.creado_por != request.user:
        messages.error(request, 'Solo puedes editar proyectos que pertenecen a tu cuenta.')
        return redirect('main:proyecto_main')

    if request.method == 'POST':
        form = ProyectoForm(request.POST, instance=proyecto)
        if form.is_valid():
            proyecto = form.save()
            messages.success(request, f'Proyecto "{proyecto.nombre}" actualizado exitosamente.')
            return redirect('main:proyecto_main')
    else:
        form = ProyectoForm(instance=proyecto)
    
    return render(request, 'esp_web/editar_proyecto.html', {
        'form': form,
        'proyecto': proyecto
    })


@login_required
@require_http_methods(["POST"])
def mover_especificacion_view(request, proyecto_id):
    """
    Vista AJAX para mover una especificación a una nueva posición
    """
    proyecto = get_object_or_404(Proyecto, id=proyecto_id, activo=True)
    
    # Verificar que el usuario es propietario del proyecto
    if proyecto.creado_por != request.user:
        return JsonResponse({'error': 'No tienes permisos para mover especificaciones de este proyecto.'}, status=403)
    
    try:
        data = json.loads(request.body)
        especificacion_id = data.get('especificacion_id')
        nueva_posicion = data.get('nueva_posicion')
        
        if not especificacion_id or not nueva_posicion:
            return JsonResponse({'error': 'Faltan datos requeridos.'}, status=400)
        
        nueva_posicion = int(nueva_posicion)
        if nueva_posicion < 1:
            return JsonResponse({'error': 'La posición debe ser mayor a 0.'}, status=400)
        
        # Obtener la especificación a mover
        especificacion = get_object_or_404(
            Especificacion,
            id=especificacion_id,
            proyecto=proyecto
        )
        
        posicion_actual = especificacion.orden
        
        if nueva_posicion == posicion_actual:
            return JsonResponse({'error': 'La nueva posición es igual a la posición actual.'}, status=400)
        
        # Obtener todas las especificaciones ordenadas
        especificaciones = list(proyecto.especificaciones.all().order_by('orden'))
        total = len(especificaciones)
        
        if nueva_posicion > total:
            return JsonResponse({'error': f'La posición máxima es {total}.'}, status=400)
        
        # Remover la especificación de su posición actual
        especificaciones.remove(especificacion)
        
        # Insertar en la nueva posición (ajustar índice porque es 0-based)
        especificaciones.insert(nueva_posicion - 1, especificacion)
        
        # Actualizar el orden de todas las especificaciones
        for orden, especificacion_item in enumerate(especificaciones, start=1):
            especificacion_item.orden = orden
            especificacion_item.save(update_fields=['orden'])
        
        return JsonResponse({
            'success': True,
            'message': f'Especificación movida de la posición {posicion_actual} a la posición {nueva_posicion}.'
        })
    
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Datos JSON inválidos.'}, status=400)
    except ValueError as e:
        return JsonResponse({'error': f'Valor inválido: {str(e)}'}, status=400)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@login_required
def obtener_imagenes_especificacion_view(request, especificacion_id):
    """
    Vista AJAX para obtener las imágenes de una especificación
    """
    especificacion = get_object_or_404(Especificacion, id=especificacion_id, proyecto__activo=True)
    
    # Verificar permisos
    if not (especificacion.proyecto.publico or especificacion.proyecto.creado_por == request.user):
        return JsonResponse({'error': 'No tienes permisos para ver las imágenes de esta especificación.'}, status=403)
    
    imagenes = especificacion.imagenes.all()
    imagenes_data = [{
        'id': img.id,
        'url': img.imagen.url if img.imagen else '',
        'descripcion': img.descripcion or '',
        'fecha_subida': img.fecha_subida.strftime('%d/%m/%Y %H:%M')
    } for img in imagenes]
    
    return JsonResponse({
        'success': True,
        'imagenes': imagenes_data
    })


@login_required
@require_http_methods(["POST"])
def subir_imagenes_especificacion_view(request, especificacion_id):
    """
    Vista AJAX para subir imágenes a una especificación
    """
    especificacion = get_object_or_404(Especificacion, id=especificacion_id, proyecto__activo=True)
    
    # Verificar que el usuario es propietario del proyecto
    if especificacion.proyecto.creado_por != request.user:
        return JsonResponse({'error': 'Solo puedes agregar imágenes a especificaciones de tus proyectos.'}, status=403)
    
    imagenes_subidas = request.FILES.getlist('imagenes')
    
    if not imagenes_subidas:
        return JsonResponse({'error': 'No se proporcionaron imágenes.'}, status=400)
    
    imagenes_creadas = []
    for imagen_file in imagenes_subidas:
        # Validar que sea una imagen
        try:
            from PIL import Image
            img = Image.open(imagen_file)
            img.verify()
            imagen_file.seek(0)  # Resetear el archivo después de verificar
        except Exception:
            continue
        
        especificacion_imagen = EspecificacionImagen(
            especificacion=especificacion,
            imagen=imagen_file
        )
        especificacion_imagen.save()
        imagenes_creadas.append({
            'id': especificacion_imagen.id,
            'url': especificacion_imagen.imagen.url
        })
    
    if not imagenes_creadas:
        return JsonResponse({'error': 'No se pudieron procesar las imágenes. Asegúrate de que sean archivos de imagen válidos.'}, status=400)
    
    return JsonResponse({
        'success': True,
        'message': f'{len(imagenes_creadas)} imagen(es) subida(s) correctamente.',
        'imagenes': imagenes_creadas
    })


@login_required
@require_http_methods(["POST"])
def eliminar_imagen_especificacion_view(request, imagen_id):
    """
    Vista AJAX para eliminar una imagen de una especificación
    """
    imagen = get_object_or_404(EspecificacionImagen, id=imagen_id)
    especificacion = imagen.especificacion
    
    # Verificar que el usuario es propietario del proyecto
    if especificacion.proyecto.creado_por != request.user:
        return JsonResponse({'error': 'Solo puedes eliminar imágenes de especificaciones de tus proyectos.'}, status=403)
    
    # Eliminar el archivo físico
    if imagen.imagen:
        imagen.imagen.delete(save=False)
    
    imagen.delete()
    
    return JsonResponse({
        'success': True,
        'message': 'Imagen eliminada correctamente.'
    })


@login_required
@require_http_methods(["POST"])
def actualizar_descripcion_imagen_view(request, imagen_id):
    """
    Vista AJAX para actualizar la descripción de una imagen
    """
    imagen = get_object_or_404(EspecificacionImagen, id=imagen_id)
    especificacion = imagen.especificacion
    
    # Verificar que el usuario es propietario del proyecto
    if especificacion.proyecto.creado_por != request.user:
        return JsonResponse({'error': 'Solo puedes editar descripciones de imágenes de tus proyectos.'}, status=403)
    
    try:
        data = json.loads(request.body)
        descripcion = data.get('descripcion', '').strip()
        
        imagen.descripcion = descripcion
        imagen.save(update_fields=['descripcion'])
        
        return JsonResponse({
            'success': True,
            'message': 'Descripción actualizada correctamente.'
        })
    
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Datos JSON inválidos.'}, status=400)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@login_required
@require_http_methods(["POST"])
def reordenar_especificaciones_view(request, proyecto_id):
    """
    Vista AJAX para reordenar las especificaciones de un proyecto
    """
    proyecto = get_object_or_404(Proyecto, id=proyecto_id, activo=True)
    
    # Verificar que el usuario es propietario del proyecto
    if proyecto.creado_por != request.user:
        return JsonResponse({'error': 'No tienes permisos para reordenar las especificaciones de este proyecto.'}, status=403)
    
    try:
        data = json.loads(request.body)
        especificaciones_ids = data.get('especificaciones_ids', [])
        
        if not especificaciones_ids:
            return JsonResponse({'error': 'No se proporcionaron IDs de especificaciones.'}, status=400)
        
        # Actualizar el orden de cada especificación
        for orden, especificacion_id in enumerate(especificaciones_ids, start=1):
            try:
                especificacion = Especificacion.objects.get(
                    id=especificacion_id,
                    proyecto=proyecto
                )
                especificacion.orden = orden
                especificacion.save(update_fields=['orden'])
            except Especificacion.DoesNotExist:
                continue
        
        return JsonResponse({'success': True, 'message': 'Orden actualizado correctamente.'})
    
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Datos JSON inválidos.'}, status=400)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

